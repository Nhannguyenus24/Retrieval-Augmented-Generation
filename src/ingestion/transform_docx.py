import re
import os
from typing import List, Dict, Tuple, Optional
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from utils.pattern import *
from dotenv import load_dotenv
from PIL import Image
import io
import json
from utils.upload_image import upload_image_to_cloudinary

# Load environment variables
load_dotenv()

# ===================== Configuration =====================
MIN_TOKENS = int(os.getenv("MIN_TOKENS", 300))
MAX_TOKENS = int(os.getenv("MAX_TOKENS", 500))
USE_TIKTOKEN = os.getenv("USE_TIKTOKEN", "false").lower() == "true"
TAB_WIDTH = int(os.getenv("TAB_WIDTH", 4))
INDENT_STEP = int(os.getenv("INDENT_STEP", 4))

# ===================== Token counter =====================
_tokenc = None

def _approx_token_count(s: str) -> int:
    return max(1, len(s.split()))

def _tok_count(s: str) -> int:
    global _tokenc
    if USE_TIKTOKEN:
        try:
            import tiktoken
            if _tokenc is None:
                _tokenc = tiktoken.get_encoding("cl100k_base")
            return len(_tokenc.encode(s))
        except Exception:
            pass
    return _approx_token_count(s)

# ===================== Image Detection & Extraction from DOCX =====================
def _extract_docx_images(docx_path: str, output_dir: str = "images") -> Tuple[List[Dict], Dict[int, List[Dict]]]:
    """
    Extract all images from DOCX file and return metadata with Cloudinary URLs
    Returns:
    - all_images: list of all image metadata
    - paragraph_images: dict mapping paragraph_index -> list of image metadata
    """
    doc = Document(docx_path)
    file_name = os.path.basename(docx_path)
    doc_id = os.path.splitext(file_name)[0]
    
    # Create output directory
    images_dir = os.path.join(output_dir, f"image_{doc_id}")
    os.makedirs(images_dir, exist_ok=True)
    
    all_images = []
    paragraph_images = {}
    image_count = 0
    
    print(f"Extracting images from DOCX: {file_name}...")
    
    # Method 1: Extract images from document.part (embedded images)
    if hasattr(doc, 'part') and hasattr(doc.part, 'related_parts'):
        for rel_id, related_part in doc.part.related_parts.items():
            if "image" in related_part.content_type:
                try:
                    image_count += 1
                    # Get image bytes
                    image_bytes = related_part.blob
                    
                    # Determine extension from content type
                    content_type = related_part.content_type
                    if 'jpeg' in content_type:
                        image_ext = 'jpg'
                    elif 'png' in content_type:
                        image_ext = 'png'
                    elif 'gif' in content_type:
                        image_ext = 'gif'
                    elif 'bmp' in content_type:
                        image_ext = 'bmp'
                    else:
                        image_ext = 'jpg'  # Default
                    
                    # Create filename
                    image_filename = f"docx_img_{image_count}.{image_ext}"
                    image_path = os.path.join(images_dir, image_filename)
                    
                    # Save image locally
                    with open(image_path, "wb") as img_file:
                        img_file.write(image_bytes)
                    
                    # Get image dimensions
                    try:
                        with Image.open(io.BytesIO(image_bytes)) as pil_img:
                            width, height = pil_img.size
                            mode = pil_img.mode
                    except:
                        width = height = mode = None
                    
                    # Upload to Cloudinary
                    cloudinary_url = ""
                    try:
                        cloudinary_url = upload_image_to_cloudinary(image_path)
                        print(f"Uploaded to Cloudinary: {cloudinary_url}")
                    except Exception as e:
                        print(f"Failed to upload {image_filename} to Cloudinary: {e}")
                    
                    image_metadata = {
                        "image_id": f"{doc_id}_img_{image_count}",
                        "filename": image_filename,
                        "file_path": os.path.abspath(image_path),
                        "cloudinary_url": cloudinary_url,
                        "image_index": image_count,
                        "format": image_ext.upper(),
                        "width": width,
                        "height": height,
                        "mode": mode,
                        "size_bytes": len(image_bytes),
                        "content_type": content_type,
                        "relationship_id": rel_id
                    }
                    
                    all_images.append(image_metadata)
                    print(f"Extracted: {image_filename} ({width}x{height}) -> {cloudinary_url}")
                    
                except Exception as e:
                    print(f"Error extracting image {image_count}: {e}")
                    continue
    
    # Method 2: Try to associate images with paragraphs containing inline shapes
    try:
        # Simple approach: distribute images evenly among paragraphs that might contain them
        total_paragraphs = len(doc.paragraphs)
        if total_paragraphs > 0 and all_images:
            print(f"Attempting to associate {len(all_images)} images with {total_paragraphs} paragraphs...")
            
            # Strategy 1: Look for paragraphs with runs that might contain images
            paragraphs_with_potential_images = []
            for para_idx, paragraph in enumerate(doc.paragraphs):
                # Check if paragraph has inline shapes or drawing elements
                has_potential_image = False
                
                # Check for runs with embedded objects
                for run in paragraph.runs:
                    if hasattr(run._element, 'drawing_lst') and run._element.drawing_lst:
                        has_potential_image = True
                        break
                    # Check for inline shapes in run element
                    if hasattr(run._element, 'xpath'):
                        try:
                            drawings = run._element.xpath('.//w:drawing')
                            if drawings:
                                has_potential_image = True
                                break
                        except:
                            pass
                
                if has_potential_image:
                    paragraphs_with_potential_images.append(para_idx)
            
            print(f"Found {len(paragraphs_with_potential_images)} paragraphs with potential images")
            
            # Strategy 2: If no specific paragraphs found, distribute images evenly
            if not paragraphs_with_potential_images and all_images:
                # Distribute images among non-empty paragraphs
                non_empty_paragraphs = []
                for para_idx, paragraph in enumerate(doc.paragraphs):
                    if paragraph.text.strip():
                        non_empty_paragraphs.append(para_idx)
                
                if non_empty_paragraphs:
                    # Distribute images across paragraphs
                    images_per_section = max(1, len(all_images) // len(non_empty_paragraphs))
                    
                    for i, img in enumerate(all_images):
                        # Assign to paragraph based on image index
                        para_idx = non_empty_paragraphs[min(i // images_per_section, len(non_empty_paragraphs) - 1)]
                        img["paragraph_index"] = para_idx
                        
                        if para_idx not in paragraph_images:
                            paragraph_images[para_idx] = []
                        paragraph_images[para_idx].append(img)
                    
                    print(f"Distributed {len(all_images)} images across {len(paragraph_images)} paragraphs")
            
            # Strategy 3: Use the paragraphs identified as having potential images
            elif paragraphs_with_potential_images:
                images_per_para = max(1, len(all_images) // len(paragraphs_with_potential_images))
                
                for i, img in enumerate(all_images):
                    para_idx = paragraphs_with_potential_images[min(i // images_per_para, len(paragraphs_with_potential_images) - 1)]
                    img["paragraph_index"] = para_idx
                    
                    if para_idx not in paragraph_images:
                        paragraph_images[para_idx] = []
                    paragraph_images[para_idx].append(img)
                
                print(f"Associated {len(all_images)} images with {len(paragraph_images)} paragraphs containing image elements")
    
    except Exception as e:
        print(f"Warning: Could not associate images with paragraphs: {e}")
        
        # Fallback: Distribute images evenly across all non-empty paragraphs
        try:
            non_empty_paragraphs = []
            for para_idx, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    non_empty_paragraphs.append(para_idx)
            
            if non_empty_paragraphs and all_images:
                images_per_section = max(1, len(all_images) // len(non_empty_paragraphs))
                
                for i, img in enumerate(all_images):
                    para_idx = non_empty_paragraphs[min(i // images_per_section, len(non_empty_paragraphs) - 1)]
                    img["paragraph_index"] = para_idx
                    
                    if para_idx not in paragraph_images:
                        paragraph_images[para_idx] = []
                    paragraph_images[para_idx].append(img)
                
                print(f"Fallback: Distributed {len(all_images)} images across {len(paragraph_images)} paragraphs")
        except Exception as fallback_error:
            print(f"Fallback failed: {fallback_error}")
            # Last resort: assign all images to first paragraph
            if all_images and doc.paragraphs:
                paragraph_images[0] = all_images
                for img in all_images:
                    img["paragraph_index"] = 0
                print(f"Last resort: Assigned all {len(all_images)} images to paragraph 0")
    
    total_images = len(all_images)
    if total_images > 0:
        print(f"Successfully extracted {total_images} images from DOCX")
        
        # Save images metadata
        metadata_file = os.path.join(images_dir, "images_metadata.json")
        with open(metadata_file, "w", encoding="utf-8") as f:
            json.dump({
                "all_images": all_images,
                "paragraph_images": paragraph_images,
                "total_count": total_images
            }, f, indent=2, ensure_ascii=False)
        print(f"Images metadata saved to {metadata_file}")
    
    return all_images, paragraph_images

def _get_image_context_docx(paragraphs: List, para_idx: int, context_range: int = 2) -> str:
    """
    Extract text context around a paragraph containing images
    """
    context_text = []
    
    # Get context from surrounding paragraphs
    start_idx = max(0, para_idx - context_range)
    end_idx = min(len(paragraphs), para_idx + context_range + 1)
    
    for i in range(start_idx, end_idx):
        if i < len(paragraphs):
            text = paragraphs[i].text.strip()
            if text and i != para_idx:  # Don't include the image paragraph itself
                context_text.append(text)
    
    return "\n\n".join(context_text)

# ===================== DOCX -> text with image detection =====================
# ===================== DOCX -> text with image detection =====================
def _docx_to_text_with_images(docx_path: str) -> Tuple[str, List[Dict], List[Dict], Dict[int, List[Dict]]]:
    """
    Extract text from DOCX with paragraph structure and image detection
    Returns:
    - doc_text: full document text with paragraph breaks
    - paragraph_info: list of dicts with paragraph metadata
    - all_images: list of all image metadata
    - paragraph_images: dict mapping paragraph_index -> list of image metadata
    """
    # First extract images
    all_images, paragraph_images = _extract_docx_images(docx_path)
    
    # Then extract text with structure
    doc = Document(docx_path)
    pieces = []
    paragraph_info = []
    
    print(f"Processing {len(doc.paragraphs)} paragraphs...")
    print(f"Images to associate: {len(all_images)}")
    print(f"Paragraph images mapping: {len(paragraph_images)} paragraphs have images")
    
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        
        # Always check for images first, even for paragraphs without text
        para_images = paragraph_images.get(i, [])
        has_images = len(para_images) > 0
        
        # If paragraph has images but no text, create placeholder text
        if has_images and not text:
            text = "[Image-only paragraph]"
        
        # Skip completely empty paragraphs (no text and no images)
        if not text and not has_images:
            continue
            
        # Get style information
        style_name = paragraph.style.name if paragraph.style else "Normal"
        
        # Detect heading level from style
        heading_level = 0
        if "Heading" in style_name:
            try:
                heading_level = int(re.search(r'\d+', style_name).group())
            except:
                heading_level = 1
        
        # Calculate indent level (approximation from left indent)
        indent_level = 0
        if paragraph.paragraph_format.left_indent:
            # Convert to points and estimate indent level
            indent_points = paragraph.paragraph_format.left_indent.pt
            indent_level = int(indent_points / 18)  # Approximate: 18pt = 1 indent level
        
        # Debug info
        if has_images:
            print(f"DEBUG: Found paragraph {i} with {len(para_images)} images")
        
        # Add image context to text if images exist
        if has_images:
            image_context = _get_image_context_docx(doc.paragraphs, i)
            for img_meta in para_images:
                # Add image reference to text
                image_ref = f"\n[IMAGE: {img_meta['filename']} - {img_meta.get('width', 'unknown')}x{img_meta.get('height', 'unknown')}]\n"
                if img_meta.get('cloudinary_url'):
                    image_ref += f"[IMAGE_URL: {img_meta['cloudinary_url']}]\n"
                if image_context:
                    image_ref += f"[IMAGE_CONTEXT: {image_context[:200]}...]\n"
                text += image_ref
        
        paragraph_info.append({
            "index": i,
            "text": text,
            "style": style_name,
            "heading_level": heading_level,
            "indent": indent_level,
            "is_heading": heading_level > 0,
            "has_images": has_images,
            "images": para_images,
            "image_count": len(para_images)
        })
        
        pieces.append(text)
    
    # Debug: Count paragraphs with images
    paragraphs_with_images = sum(1 for p in paragraph_info if p['has_images'])
    total_images_in_paragraphs = sum(p['image_count'] for p in paragraph_info)
    
    print(f"Final result: {paragraphs_with_images} paragraphs with images")
    print(f"Total images in paragraphs: {total_images_in_paragraphs}")
    
    return "\n\n".join(pieces), paragraph_info, all_images, paragraph_images

def _docx_to_text_with_structure(docx_path: str) -> Tuple[str, List[Dict]]:
    """
    Extract text from DOCX with paragraph structure preserved
    Returns:
    - doc_text: full document text with paragraph breaks
    - paragraph_info: list of dicts with paragraph metadata
    """
    doc = Document(docx_path)
    pieces = []
    paragraph_info = []
    
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
            
        # Get style information
        style_name = paragraph.style.name if paragraph.style else "Normal"
        
        # Detect heading level from style
        heading_level = 0
        if "Heading" in style_name:
            try:
                heading_level = int(re.search(r'\d+', style_name).group())
            except:
                heading_level = 1
        
        # Calculate indent level (approximation from left indent)
        indent_level = 0
        if paragraph.paragraph_format.left_indent:
            # Convert to points and estimate indent level
            indent_points = paragraph.paragraph_format.left_indent.pt
            indent_level = int(indent_points / 18)  # Approximate: 18pt = 1 indent level
        
        paragraph_info.append({
            "index": i,
            "text": text,
            "style": style_name,
            "heading_level": heading_level,
            "indent": indent_level,
            "is_heading": heading_level > 0
        })
        
        pieces.append(text)
    
    return "\n\n".join(pieces), paragraph_info

# ===================== Heading detection for DOCX =====================
def _is_docx_heading(para_info: Dict) -> Tuple[bool, int]:
    """
    Detect if a paragraph is a heading based on DOCX metadata
    Returns: (is_heading, level)
    """
    # First check style-based headings
    if para_info["is_heading"] and para_info["heading_level"] > 0:
        return True, para_info["heading_level"]
    
    # Fallback to pattern-based detection
    text = para_info["text"].strip()
    if not text or len(text) < 3:
        return False, 0
    
    # Check numbered headings (1., 1.1., 1.1.1.)
    numbered_match = re.match(r"^\s*(\d+\.)+\s+(.+)$", text)
    if numbered_match:
        dots = numbered_match.group(1).count('.')
        return True, min(dots, 6)
    
    # Check other patterns from utils.pattern
    for i, pattern in enumerate(HEADING_PATTERNS):
        if pattern.match(text):
            if i == 0:  # numbered pattern already handled
                continue
            # ALL CAPS = level 1, others = level 2-3
            level = 1 if i == 1 else (2 if len(text) < 50 else 3)
            return True, level
    
    return False, 0

def _extract_headings_from_paragraphs(paragraph_info: List[Dict]) -> List[Dict]:
    """
    Extract headings from paragraph info and create heading hierarchy
    """
    headings = []
    
    for para in paragraph_info:
        is_head, level = _is_docx_heading(para)
        if is_head:
            headings.append({
                "text": para["text"],
                "level": level,
                "paragraph_index": para["index"],
                "indent": para["indent"]
            })
    
    return headings

def _build_breadcrumbs_docx(headings: List[Dict], current_para_index: int) -> List[str]:
    """
    Build breadcrumb path for a given paragraph based on preceding headings
    """
    breadcrumbs = []
    current_levels = {}  # level -> heading text
    
    # Find headings that precede current paragraph
    relevant_headings = [h for h in headings if h["paragraph_index"] < current_para_index]
    
    for heading in relevant_headings:
        level = heading["level"]
        text = heading["text"]
        
        # Clear deeper levels when we encounter a higher level heading
        keys_to_remove = [k for k in current_levels.keys() if k >= level]
        for k in keys_to_remove:
            del current_levels[k]
        
        # Set current level
        current_levels[level] = text
    
    # Build breadcrumb path from level 1 to deepest
    if current_levels:
        for level in sorted(current_levels.keys()):
            breadcrumbs.append(current_levels[level])
    
    return breadcrumbs

# ===================== Convert paragraphs to blocks =====================
def _paragraphs_to_blocks(paragraph_info: List[Dict]) -> List[Dict]:
    """
    Convert paragraph info into blocks based on indent and content similarity
    """
    if not paragraph_info:
        return []
    
    blocks = []
    current_block = []
    current_indent = paragraph_info[0]["indent"] if paragraph_info else 0
    
    def flush_block():
        nonlocal current_block, current_indent
        if current_block:
            combined_text = "\n".join([p["text"] for p in current_block])
            blocks.append({
                "indent": current_indent,
                "text": combined_text,
                "paragraph_indices": [p["index"] for p in current_block]
            })
            current_block = []
    
    for para in paragraph_info:
        text = para["text"]
        indent = para["indent"]
        
        # Skip empty paragraphs
        if not text.strip():
            continue
        
        # If indent changes significantly, flush current block
        if abs(indent - current_indent) >= INDENT_STEP:
            flush_block()
            current_indent = indent
            current_block = [para]
        else:
            current_block.append(para)
    
    # Flush remaining block
    flush_block()
    return blocks

# ===================== Split large blocks by sentences =====================
def _split_docx_block_into_chunks(blk_text: str, min_tokens=MIN_TOKENS, max_tokens=MAX_TOKENS) -> List[str]:
    """
    Split a DOCX block into chunks if it's too large
    """
    # Split by paragraphs first, then by sentences if needed
    paragraphs = blk_text.split('\n')
    parts: List[str] = []
    
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        
        # Check if it's a bullet point
        if BULLET_PATTERN.match(para):
            parts.append(para)
        else:
            # Split by sentences for regular paragraphs
            segs = SPLIT_PATTERN.split(para)
            for s in segs:
                s = s.strip()
                if s:
                    parts.append(s)
    
    # Now group parts into chunks
    chunks: List[str] = []
    cur: List[str] = []
    cur_tok = 0
    
    for p in parts:
        t = _tok_count(p)
        if t > max_tokens:
            # Single part too long: push separately
            if cur:
                if cur_tok >= min_tokens:
                    chunks.append("\n".join(cur))
                    cur, cur_tok = [], 0
                else:
                    chunks.append("\n".join(cur + [p]))
                    cur, cur_tok = [], 0
            else:
                chunks.append(p)
            continue
        
        if cur and cur_tok + t > max_tokens:
            chunks.append("\n".join(cur))
            cur, cur_tok = [p], t
        else:
            cur.append(p)
            cur_tok += t
    
    if cur:
        if chunks and _tok_count(chunks[-1]) + cur_tok <= max_tokens:
            chunks[-1] = (chunks[-1] + "\n" + "\n".join(cur)).strip("\n")
        else:
            chunks.append("\n".join(cur))
    
    # Merge small final chunk if possible
    if len(chunks) >= 2 and _tok_count(chunks[-1]) < min_tokens:
        if _tok_count(chunks[-2]) + _tok_count(chunks[-1]) <= max_tokens:
            chunks[-2] = (chunks[-2] + "\n" + chunks[-1]).strip("\n")
            chunks.pop()
    
    return chunks

# ===================== Pack blocks into chunks =====================
def _pack_docx_blocks_into_chunks(blocks: List[Dict], headings: List[Dict], min_tokens=MIN_TOKENS, max_tokens=MAX_TOKENS) -> List[Dict]:
    """
    Group DOCX blocks into ~500 token chunks
    """
    out: List[Dict] = []
    cur_parts: List[str] = []
    cur_inds: List[int] = []
    cur_block_indices: List[int] = []
    cur_tok = 0
    
    def close():
        nonlocal cur_parts, cur_inds, cur_block_indices, cur_tok
        if cur_parts:
            payload = "\n".join(cur_parts).strip("\n")
            # Get breadcrumbs for the first block in this chunk
            first_block_idx = cur_block_indices[0] if cur_block_indices else 0
            first_para_idx = blocks[first_block_idx]["paragraph_indices"][0] if first_block_idx < len(blocks) and blocks[first_block_idx]["paragraph_indices"] else 0
            breadcrumbs = _build_breadcrumbs_docx(headings, first_para_idx)
            
            out.append({
                "payload": payload,
                "indent_levels": sorted(set(cur_inds)),
                "breadcrumbs": breadcrumbs,
                "block_indices": cur_block_indices.copy()
            })
            cur_parts = []
            cur_inds = []
            cur_block_indices = []
            cur_tok = 0
    
    for block_idx, blk in enumerate(blocks):
        txt = blk["text"]
        ind = blk["indent"]
        t = _tok_count(txt)
        
        # If block is too large, split it
        if t > max_tokens:
            subchunks = _split_docx_block_into_chunks(txt, min_tokens, max_tokens)
            for sc in subchunks:
                st = _tok_count(sc)
                if cur_tok > 0 and cur_tok + st <= max_tokens:
                    cur_parts.append(sc)
                    cur_inds.append(ind)
                    if block_idx not in cur_block_indices:
                        cur_block_indices.append(block_idx)
                    cur_tok += st
                    if cur_tok >= min_tokens:
                        close()
                else:
                    if cur_tok > 0:
                        close()
                    cur_parts.append(sc)
                    cur_inds.append(ind)
                    cur_block_indices.append(block_idx)
                    cur_tok = st
                    if cur_tok >= min_tokens:
                        close()
            continue
        
        # Regular sized block
        if cur_tok + t <= max_tokens:
            cur_parts.append(txt)
            cur_inds.append(ind)
            cur_block_indices.append(block_idx)
            cur_tok += t
        else:
            # Close current chunk and start new one
            close()
            cur_parts.append(txt)
            cur_inds.append(ind)
            cur_block_indices.append(block_idx)
            cur_tok = t
            if cur_tok >= min_tokens:
                close()
    
    close()
    return out

# ===================== Enhanced chunk creation with image info for DOCX =====================
def _pack_docx_blocks_with_images(blocks: List[Dict], headings: List[Dict], paragraph_images: Dict[int, List[Dict]], 
                                 min_tokens=MIN_TOKENS, max_tokens=MAX_TOKENS) -> List[Dict]:
    """
    Enhanced version that includes image information in chunks for DOCX
    """
    out: List[Dict] = []
    cur_parts: List[str] = []
    cur_inds: List[int] = []
    cur_block_indices: List[int] = []
    cur_tok = 0
    
    def close():
        nonlocal cur_parts, cur_inds, cur_block_indices, cur_tok
        if cur_parts:
            payload = "\n".join(cur_parts).strip("\n")
            # Get breadcrumbs for the first block in this chunk
            first_block_idx = cur_block_indices[0] if cur_block_indices else 0
            first_para_idx = blocks[first_block_idx]["paragraph_indices"][0] if first_block_idx < len(blocks) and blocks[first_block_idx]["paragraph_indices"] else 0
            breadcrumbs = _build_breadcrumbs_docx(headings, first_para_idx)
            
            # Collect images for this chunk
            chunk_images = []
            chunk_image_urls = []
            
            # Get all paragraph indices covered by this chunk
            all_para_indices = []
            for block_idx in cur_block_indices:
                if block_idx < len(blocks):
                    all_para_indices.extend(blocks[block_idx]["paragraph_indices"])
            
            # Collect images from all paragraphs in this chunk
            for para_idx in all_para_indices:
                if para_idx in paragraph_images:
                    for img in paragraph_images[para_idx]:
                        chunk_images.append(img)
                        if img.get('cloudinary_url'):
                            chunk_image_urls.append(img['cloudinary_url'])
            
            out.append({
                "payload": payload,
                "indent_levels": sorted(set(cur_inds)),
                "breadcrumbs": breadcrumbs,
                "block_indices": cur_block_indices.copy(),
                "images": chunk_images,  # NEW: Include images in chunk
                "image_urls": chunk_image_urls,  # NEW: Cloudinary URLs array
                "has_images": len(chunk_images) > 0  # NEW: Flag for quick check
            })
            cur_parts = []
            cur_inds = []
            cur_block_indices = []
            cur_tok = 0
    
    # Rest of the logic remains the same as _pack_docx_blocks_into_chunks
    for block_idx, blk in enumerate(blocks):
        txt = blk["text"]
        ind = blk["indent"]
        t = _tok_count(txt)
        
        # If block is too large, split it
        if t > max_tokens:
            subchunks = _split_docx_block_into_chunks(txt, min_tokens, max_tokens)
            for sc in subchunks:
                st = _tok_count(sc)
                if cur_tok > 0 and cur_tok + st <= max_tokens:
                    cur_parts.append(sc)
                    cur_inds.append(ind)
                    if block_idx not in cur_block_indices:
                        cur_block_indices.append(block_idx)
                    cur_tok += st
                    if cur_tok >= min_tokens:
                        close()
                else:
                    if cur_tok > 0:
                        close()
                    cur_parts.append(sc)
                    cur_inds.append(ind)
                    cur_block_indices.append(block_idx)
                    cur_tok = st
                    if cur_tok >= min_tokens:
                        close()
            continue
        
        # Regular sized block
        if cur_tok + t <= max_tokens:
            cur_parts.append(txt)
            cur_inds.append(ind)
            cur_block_indices.append(block_idx)
            cur_tok += t
        else:
            # Close current chunk and start new one
            close()
            cur_parts.append(txt)
            cur_inds.append(ind)
            cur_block_indices.append(block_idx)
            cur_tok = t
            if cur_tok >= min_tokens:
                close()
    
    close()
    return out

# ===================== Main API =====================
def docx_to_chunks_smart_indent(
    docx_path: str,
    min_tokens: int = MIN_TOKENS,
    max_tokens: int = MAX_TOKENS,
    extract_images: bool = True  # NEW parameter
) -> List[Dict]:
    """
    Enhanced version with image detection and extraction
    - Read DOCX -> text (preserve paragraph structure)
    - Divide into blocks by indent and content similarity
    - Pack blocks into ~500 token chunks (within [300, 500] if possible)
    - Automatically detect headings and create breadcrumbs
    - Prepend heading context to the chunk content
    - NEW: Detect and extract images, upload to Cloudinary
    """
    file_name = os.path.basename(docx_path)
    doc_id = os.path.splitext(file_name)[0]
    
    if extract_images:
        # Use enhanced text extraction with image detection
        doc_text, paragraph_info, all_images, paragraph_images = _docx_to_text_with_images(docx_path)
        blocks = _paragraphs_to_blocks(paragraph_info)
        headings = _extract_headings_from_paragraphs(paragraph_info)
        
        # Use enhanced packing with image info
        packed = _pack_docx_blocks_with_images(blocks, headings, paragraph_images, min_tokens, max_tokens)
    else:
        # Original logic without images
        doc_text, paragraph_info = _docx_to_text_with_structure(docx_path)
        blocks = _paragraphs_to_blocks(paragraph_info)
        headings = _extract_headings_from_paragraphs(paragraph_info)
        packed = _pack_docx_blocks_into_chunks(blocks, headings, min_tokens, max_tokens)
    
    results = []
    
    for idx, ch in enumerate(packed):
        # Build breadcrumb string
        breadcrumb_text = ""
        if ch["breadcrumbs"]:
            breadcrumb_text = " > ".join(ch["breadcrumbs"])
        
        # Prepend breadcrumbs to content if available
        content_with_context = ch["payload"]
        if breadcrumb_text:
            content_with_context = f"[Context: {breadcrumb_text}]\n\n{ch['payload']}"
        
        # Enhanced metadata with image information
        chunk_data = {
            "chunk_id": idx,
            "file_name": file_name,
            "indent_levels": ch["indent_levels"],
            "content": content_with_context,
            "token_est": _tok_count(content_with_context),
            "doc_id": doc_id,
            "source": os.path.abspath(docx_path),
            "section_title": ch["breadcrumbs"][-1] if ch["breadcrumbs"] else None,
            "heading_path": breadcrumb_text if breadcrumb_text else None,
            "chunk_index": idx,
            "breadcrumbs": ch["breadcrumbs"]
        }
        
        # Add image-related metadata if available
        if extract_images and ch.get("has_images"):
            chunk_data.update({
                "has_images": ch["has_images"],
                "images": ch["images"],
                "image_urls": ch["image_urls"],  # Array of Cloudinary URLs
                "image_count": len(ch["images"]),
                "image_filenames": [img["filename"] for img in ch["images"]]
            })
        else:
            chunk_data.update({
                "has_images": False,
                "images": [],
                "image_urls": [],
                "image_count": 0,
                "image_filenames": []
            })
        
        results.append(chunk_data)
    
    # Print summary
    if extract_images:
        total_chunks = len(results)
        chunks_with_images = sum(1 for r in results if r["has_images"])
        total_images = sum(r["image_count"] for r in results)
        
        print(f"\nDOCX CHUNKING SUMMARY:")
        print(f"Total chunks: {total_chunks}")
        print(f"Chunks with images: {chunks_with_images}")
        print(f"Total images detected: {total_images}")
    
    return results

# ===================== Utility function to get DOCX image info =====================
def get_docx_chunk_images_info(chunks: List[Dict]) -> Dict:
    """
    Get summary of images across all DOCX chunks
    """
    image_summary = {
        "total_chunks": len(chunks),
        "chunks_with_images": 0,
        "total_images": 0,
        "image_formats": {},
        "large_images": [],  # Images > 1MB
        "cloudinary_urls": []  # All Cloudinary URLs
    }
    
    for chunk in chunks:
        if chunk.get("has_images"):
            image_summary["chunks_with_images"] += 1
            
            # Add Cloudinary URLs
            image_summary["cloudinary_urls"].extend(chunk.get("image_urls", []))
            
            for img in chunk.get("images", []):
                image_summary["total_images"] += 1
                
                # Count by format
                fmt = img.get("format", "UNKNOWN")
                if fmt not in image_summary["image_formats"]:
                    image_summary["image_formats"][fmt] = 0
                image_summary["image_formats"][fmt] += 1
                
                # Track large images
                if img.get("size_bytes", 0) > 1024 * 1024:  # > 1MB
                    image_summary["large_images"].append({
                        "filename": img["filename"],
                        "cloudinary_url": img.get("cloudinary_url", ""),
                        "size_mb": round(img["size_bytes"] / (1024 * 1024), 2),
                        "dimensions": f"{img.get('width', 'unknown')}x{img.get('height', 'unknown')}"
                    })
    
    return image_summary
